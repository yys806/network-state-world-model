"""Build the 2026-05-31 internship summary DOCX."""

from __future__ import annotations

import json
import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = ROOT.parent
OUTPUT_DIR = WORKSPACE_ROOT / "文档" / "组会" / "5.31"
FIGURE_DIR = OUTPUT_DIR / "figures"
DOCX_PATH = OUTPUT_DIR / "科研实习进组以来工作总结_禹尧珅_20260531.docx"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_rounded_rectangle(draw: ImageDraw.ImageDraw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def make_timeline_figure(path: Path) -> None:
    width, height = 1800, 920
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_f = font(54, True)
    h_f = font(34, True)
    body_f = font(25)
    small_f = font(22)

    draw.text((70, 45), "进组以来研究主线：从问题理解到 PI-JWM 双图世界模型", fill="#14213d", font=title_f)
    draw.text((74, 120), "AirFogSim 作为仿真与数据工具，主线是物理网络-信息网络联合世界模型", fill="#536271", font=body_f)

    stages = [
        ("1 课题理解", "ST-GNN / World Model / 低空网络\n明确物理-信息联合预测问题"),
        ("2 仿真机制", "解析任务产生、信道衰减、资源调度\n形成可追踪的日志与样本"),
        ("3 动作条件", "dataset / strict action / edge action\n从状态预测走向动作条件预测"),
        ("4 世界模型", "v0-v3: 统一接口、latent rollout\n通信图消息传递与诊断"),
        ("5 双图 v6", "物理图 + 信息图 + 动作历史\n预测节点、链路、任务未来状态"),
    ]
    colors = ["#E8F2FF", "#EAF7EE", "#FFF4DF", "#F1ECFF", "#FFECEC"]
    outline = ["#4D8DD8", "#53A66D", "#D79B35", "#8064C8", "#D86B6B"]
    x0, y0, gap = 70, 230, 28
    card_w, card_h = 315, 470
    for i, (stage, text) in enumerate(stages):
        x = x0 + i * (card_w + gap)
        draw_rounded_rectangle(draw, (x, y0, x + card_w, y0 + card_h), 26, colors[i], outline[i], 3)
        draw.text((x + 26, y0 + 28), stage, fill="#1f2933", font=h_f)
        draw.line((x + 28, y0 + 86, x + card_w - 28, y0 + 86), fill=outline[i], width=4)
        draw.multiline_text((x + 26, y0 + 120), text, fill="#263238", font=body_f, spacing=12)
        if i < len(stages) - 1:
            ax = x + card_w + 4
            ay = y0 + card_h // 2
            draw.line((ax, ay, ax + gap - 8, ay), fill="#6b7280", width=4)
            draw.polygon([(ax + gap - 8, ay - 10), (ax + gap + 8, ay), (ax + gap - 8, ay + 10)], fill="#6b7280")

    draw_rounded_rectangle(draw, (155, 765, 1645, 850), 24, "#F6F8FA", "#CBD5E1", 2)
    draw.text(
        (190, 788),
        "当前阶段结论：v6 已经能训练、评估并形成初步消融结论；下一步补普通模型对比和 active-rate 改进。",
        fill="#1f2933",
        font=small_f,
    )
    img.save(path)


def make_v6_metric_figure(summary: dict, path: Path) -> None:
    runs = summary["real_data_sanity"]["runs"]
    modes = ["dual", "physical_only", "information_only"]
    labels = {"dual": "dual 双图", "physical_only": "physical", "information_only": "information"}
    active = [runs[m]["test_eval"]["active_rate"]["active_rmse"] for m in modes]
    link = [runs[m]["test_eval"]["link_rate"]["rmse"] for m in modes]
    task = [runs[m]["test_eval"]["task"]["rmse"] for m in modes]

    width, height = 1700, 980
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_f = font(52, True)
    axis_f = font(26, True)
    txt_f = font(24)
    small_f = font(22)

    draw.text((70, 42), "PI-JWM v6 full80 消融结果：双图对活跃链路速率更有帮助", fill="#14213d", font=title_f)
    draw.text((72, 112), "测试集指标，三种模式 activity F1 均为 1.0；下图重点比较误差项。", fill="#536271", font=txt_f)

    plot_x, plot_y = 120, 210
    plot_w, plot_h = 1460, 540
    max_v = max(active) * 1.15
    colors = ["#2563EB", "#16A34A", "#F97316"]
    group_w = plot_w / 3
    bar_w = 110
    for i, mode in enumerate(modes):
        cx = plot_x + group_w * i + group_w / 2
        h = active[i] / max_v * plot_h
        x1 = cx - bar_w / 2
        y1 = plot_y + plot_h - h
        draw.rounded_rectangle((x1, y1, x1 + bar_w, plot_y + plot_h), radius=14, fill=colors[i])
        draw.text((x1 - 22, y1 - 40), f"{active[i]:.1f}", fill="#111827", font=txt_f)
        draw.text((cx - 86, plot_y + plot_h + 25), labels[mode], fill="#111827", font=axis_f)
    draw.line((plot_x, plot_y + plot_h, plot_x + plot_w, plot_y + plot_h), fill="#374151", width=3)
    draw.text((plot_x, plot_y + plot_h + 78), "active-rate RMSE 越低越好", fill="#536271", font=small_f)

    table_x, table_y = 150, 820
    col_w = [350, 300, 300, 300]
    headers = ["模式", "active-rate RMSE", "link-rate RMSE", "task RMSE"]
    row_h = 46
    draw_rounded_rectangle(draw, (table_x - 20, table_y - 20, table_x + sum(col_w) + 20, table_y + row_h * 4 + 25), 18, "#F8FAFC", "#CBD5E1", 2)
    x = table_x
    for j, header in enumerate(headers):
        draw.text((x + 10, table_y), header, fill="#111827", font=axis_f)
        x += col_w[j]
    for i, mode in enumerate(modes):
        y = table_y + row_h * (i + 1)
        values = [labels[mode], f"{active[i]:.3f}", f"{link[i]:.3f}", f"{task[i]:.3f}"]
        x = table_x
        for j, value in enumerate(values):
            fill = "#111827" if not (mode == "dual" and j in (1, 2)) else "#1D4ED8"
            draw.text((x + 10, y), value, fill=fill, font=txt_f)
            x += col_w[j]
    img.save(path)


def add_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2EC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        add_border(p)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_docx() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    timeline_path = FIGURE_DIR / "research_timeline.png"
    metrics_path = FIGURE_DIR / "v6_metrics_comparison.png"
    mechanism_path = ROOT / "artifacts" / "experiments" / "airfogsim_v0" / "figures" / "airfogsim_state_transition_flow.png"
    v4_path = ROOT / "artifacts" / "experiments" / "airfogsim_v0" / "figures" / "world_model_v4_dual_graph_ablation_compare.png"
    summary_path = ROOT / "artifacts" / "experiments" / "pi_jwm_v6_eval_full80" / "v6_dual_graph_smoke_summary.json"
    summary = json.loads(summary_path.read_text(encoding="utf-8"))

    make_timeline_figure(timeline_path)
    make_v6_metric_figure(summary, metrics_path)

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Microsoft YaHei"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("科研实习进组以来工作总结")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(20, 33, 61)
    subtitle = doc.add_paragraph("禹尧珅｜2026 年 5 月 31 日线上讨论")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "一、15 分钟汇报摘要", 1)
    doc.add_paragraph(
        "进组以来，我围绕“基座模型驱动的网络状态联合预测”这一方向展开工作。"
        "最初的任务是理解网联具身智能体协同场景中物理网络和信息网络为何需要联合预测，随后将问题收缩到"
        "无人机/车辆/边缘节点协同下的通信资源、计算卸载与运动状态联合演化。"
        "当前形成的主线是 PI-JWM，即 Physical-Information Joint World Model。"
        "AirFogSim/SUMO 在本工作中作为仿真器和数据生成工具，而不是研究对象本身。"
    )
    doc.add_picture(str(timeline_path), width=Cm(16.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "二、问题 Formulation", 1)
    doc.add_paragraph(
        "老师给出的具体问题可以概括为：在网络基站覆盖环境下，如何对无人机群的通信资源、计算卸载与运动轨迹进行在线协同决策，"
        "使任务时延、能耗同时优化，并在链路波动、任务随机到达和局部观测下保持系统稳定。"
        "这要求模型不仅预测单一流量或单条链路，而是预测物理侧和信息侧的协同演化。"
    )
    add_bullets(
        doc,
        [
            "状态：节点位置、速度、队列、任务、链路活动、链路速率、资源使用情况。",
            "动作：卸载动作、RB/CPU 分配、任务回传、边级调度动作，以及后续可能接入的轨迹动作。",
            "目标：学习近似转移函数，预测未来节点状态、链路活跃性、链路速率和任务状态，为后续候选动作评估提供依据。",
            "核心挑战：物理移动改变网络拓扑，通信和计算决策反过来改变队列、任务完成和未来资源状态。两侧不能割裂建模。",
        ],
    )

    add_heading(doc, "三、用了什么方法", 1)
    add_heading(doc, "1. 仿真机制解析与数据构建", 2)
    doc.add_paragraph(
        "首先解析 AirFogSim 的单步推进逻辑，包括交通更新、任务生成、动作注入、无线链路速率计算、计算执行、队列变化和能耗更新。"
        "在此基础上，构建 dataset、多 seed 数据、strict action、edge action 和 world model dataset，使后续模型可以从状态预测推进到动作条件预测。"
    )
    if mechanism_path.exists():
        doc.add_picture(str(mechanism_path), width=Cm(15.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "2. 从普通预测模型到动作条件世界模型", 2)
    add_bullets(
        doc,
        [
            "state-only baseline：只看历史状态，验证最基本的预测可行性。",
            "state-action baseline：加入调度动作，判断动作是否能解释未来变化。",
            "edge-level / two-stage link model：将链路活动和链路速率拆开，定位链路侧瓶颈。",
            "world model v0-v3：从统一接口、分阶段训练、latent rollout 推进到通信图消息传递。",
            "v6 双图世界模型：用物理图、信息图和动作历史联合预测未来节点、链路和任务状态。",
        ],
    )
    if v4_path.exists():
        doc.add_picture(str(v4_path), width=Cm(15.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "四、克服了什么困难", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(["困难", "具体表现", "解决方式"]):
        hdr[i].text = text
        set_cell_shading(hdr[i], "D9EAF7")
    rows = [
        ("问题太大", "最初课题涉及基座模型、世界模型、网络状态预测、协同决策，范围很宽。", "先收缩到小规模 AirFogSim 场景，再围绕物理-信息联合预测建立主线。"),
        ("仿真机制不透明", "任务、信道、动作、资源更新的顺序不清楚，难以解释数据来源。", "从代码和日志解析 Fsim 单步推进机制，明确数据如何产生。"),
        ("动作和状态难对齐", "全局动作无法直接解释某条候选通信边未来是否活跃。", "构建 strict action 和 edge action，把调度动作映射到边级特征。"),
        ("链路标签极稀疏", "大多数候选边不活跃，直接回归速率容易被零值主导。", "拆成 activity prediction 和 active-rate regression，并补充 F1、active-rate RMSE 等指标。"),
        ("容易偏向 selector", "v5 决策排序容易把主线带到动作选择接口上。", "重新明确 PI-JWM 主线，v5 只作为决策接口诊断，v6 回到双图状态预测。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text

    add_heading(doc, "五、有什么创新", 1)
    add_bullets(
        doc,
        [
            "问题层面：把课题明确为物理网络-信息网络联合世界模型，而不是单一链路预测或单一流量预测。",
            "数据层面：从仿真日志中抽取节点、链路、任务、动作，并进一步构造边级动作，使动作条件预测可训练。",
            "模型层面：从普通预测模型推进到动作条件 latent rollout，再推进到双图联合 world model。",
            "评估层面：不只看总 loss，而是分别评估 activity F1、active-rate RMSE、link-rate RMSE、node RMSE、task RMSE。",
            "方法意义：模型可以作为后续通信资源、计算卸载和轨迹动作候选评估的快速近似器。",
        ],
    )

    add_heading(doc, "六、当前实验结果", 1)
    doc.add_paragraph(
        "最新 full80 GPU 评估中，数据切分为 train=1520、val=190、test=190；三种模式各训练 80 epoch。"
        "三种模式的 link activity F1 均达到 1.0，说明链路是否活跃已经可以稳定预测。"
        "真正区分模型的是 active-rate RMSE 和 link-rate RMSE。"
    )
    doc.add_picture(str(metrics_path), width=Cm(16.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    result_table = doc.add_table(rows=1, cols=6)
    result_table.style = "Table Grid"
    headers = ["模式", "activity F1", "active-rate RMSE", "link-rate RMSE", "node RMSE", "task RMSE"]
    for i, h in enumerate(headers):
        result_table.rows[0].cells[i].text = h
        set_cell_shading(result_table.rows[0].cells[i], "D9EAF7")
    values = [
        ("dual", "1.000000", "228.318431", "6.416102", "38.282549", "3.664203"),
        ("physical_only", "1.000000", "268.036374", "7.473947", "37.067164", "4.855683"),
        ("information_only", "1.000000", "235.526004", "6.563317", "37.305306", "3.380836"),
    ]
    for row in values:
        cells = result_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph(
        "阶段性结论：信息图足以判断链路是否活跃；双图联合在活跃链路速率预测和整体链路速率预测上最好；"
        "任务状态预测更依赖信息图和动作历史；物理图对节点状态预测有帮助，但单独使用不足以解释通信任务演化。"
    )

    add_heading(doc, "七、后续计划", 1)
    add_bullets(
        doc,
        [
            "补普通模型与世界模型统一对比：将 MLP、Ridge、state-action baseline、non-rollout predictor 与 PI-JWM v6 放在同一指标体系下比较。",
            "改进 active-rate 建模：当前速率幅值仍是主要误差来源，下一步尝试两阶段 activity + active-rate head、加权损失或 active-only regression。",
            "补稳健性和泛化实验：继续做跨 seed、输入扰动、阈值迁移、置信区间和 per-seed 统计。",
            "接入决策接口：在状态预测更稳定之后，用 world model 评估候选卸载、资源分配和轨迹动作。",
            "整理阶段材料：将 v6 结果、普通模型对比和后续改进整理进研究进展文档和组会汇报材料。",
        ],
    )

    add_heading(doc, "八、15 分钟讲述建议", 1)
    add_bullets(
        doc,
        [
            "2 分钟：讲问题 formulation，说明为什么需要物理-信息联合预测。",
            "3 分钟：讲仿真机制和数据构建，说明数据从哪里来。",
            "4 分钟：讲方法演进，从 baseline 到动作条件 world model，再到 v6 双图。",
            "3 分钟：讲困难和创新，重点讲动作对齐、稀疏链路、双图联合。",
            "2 分钟：讲实验结果，强调 dual 对 active-rate 的提升。",
            "1 分钟：讲下一步计划，补普通模型对比和 active-rate 改进。",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "附录：已完成工作清单", 1)
    add_bullets(
        doc,
        [
            "完成 AirFogSim/Fsim 机制解析和仿真日志整理。",
            "完成 dataset_v0、多 seed、strict action、edge action 和 world model dataset 构建。",
            "完成 state-only、state-action、structured、edge-level、two-stage、edge-action baseline。",
            "完成 world model v0/v1/v2/v3/v4/v6 的阶段性推进。",
            "完成 v6 full80 三模式消融和原尺度指标评估。",
            "完成本地 PI-JWM 项目结构整理，代码、脚本、测试和实验输出已归位。",
        ],
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
