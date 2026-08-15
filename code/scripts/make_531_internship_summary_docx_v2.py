"""Build the polished 2026-05-31 internship summary DOCX and speaker notes.

This version keeps PI-JWM as the main research line. AirFogSim/SUMO is only
described as a simulator and data source.
"""

from __future__ import annotations

import json
import math
import textwrap
from pathlib import Path

import requests
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFilter, ImageFont


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = ROOT.parent
OUTPUT_DIR = WORKSPACE_ROOT / "文档" / "组会" / "5.31"
FIGURE_DIR = OUTPUT_DIR / "figures"
DATA_SUMMARY_PATH = ROOT / "artifacts" / "experiments" / "pi_jwm_v6_eval_full80" / "v6_dual_graph_smoke_summary.json"

DOCX_PATH = OUTPUT_DIR / "科研实习进组以来工作总结_禹尧珅_20260531_v2.docx"
SCRIPT_PATH = OUTPUT_DIR / "科研实习进组以来工作总结_禹尧珅_20260531_讲稿.md"

AI_FIGURES = {
    "cover_pijwm_scene.png": (
        "clean technical isometric illustration of connected UAVs, autonomous vehicles, roadside edge servers "
        "and base stations in a modern city grid, wireless communication links and mobility paths as two coordinated "
        "layers, academic research report style, white background, deep blue teal orange accents, no readable text, "
        "no letters, no watermark, no logo"
    ),
    "problem_formulation_ai.png": (
        "scientific illustration for problem formulation, current connected network state on the left, action "
        "intervention in the middle, future network state on the right, drones vehicles edge nodes wireless links "
        "and task flows, clean 3d infographic, no readable text, no letters, no watermark, no logo"
    ),
    "method_pipeline_ai.png": (
        "professional abstract workflow image, simulator logs and datasets flowing into neural model training, "
        "then rollout prediction and diagnostics, data blocks graph modules and evaluation charts, horizontal "
        "pipeline, clean academic 3d infographic, no readable text, no letters, no watermark, no logo"
    ),
    "dual_graph_model_ai.png": (
        "Physical Information Joint World Model concept without text, two stacked graph layers converging into a "
        "central prediction model cube, physical mobility graph in blue, information task communication graph in "
        "teal green, future rollout outputs in orange, clean white background, no readable text, no watermark"
    ),
}


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def wrap_text(text: str, width: int) -> str:
    lines: list[str] = []
    for raw in text.splitlines():
        if not raw:
            lines.append("")
            continue
        line = ""
        count = 0
        for ch in raw:
            step = 2 if ord(ch) > 127 else 1
            if count + step > width:
                lines.append(line)
                line = ch
                count = step
            else:
                line += ch
                count += step
        if line:
            lines.append(line)
    return "\n".join(lines)


def rounded(draw: ImageDraw.ImageDraw, box, radius: int, fill, outline=None, width: int = 1) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def add_shadow(base: Image.Image, box, radius: int = 28, offset=(0, 12), blur: int = 18) -> None:
    shadow = Image.new("RGBA", base.size, (0, 0, 0, 0))
    sd = ImageDraw.Draw(shadow)
    shifted = (box[0] + offset[0], box[1] + offset[1], box[2] + offset[0], box[3] + offset[1])
    sd.rounded_rectangle(shifted, radius=radius, fill=(15, 23, 42, 32))
    shadow = shadow.filter(ImageFilter.GaussianBlur(blur))
    base.alpha_composite(shadow)


def fetch_ai_image(filename: str, prompt: str, width: int = 1536, height: int = 864) -> Path:
    """Fetch an AI-style bitmap from a public image endpoint.

    The preferred built-in image generation tool is not exposed in this local
    session, so this function uses a network fallback that does not require a
    project API key. Generated figures are still treated as AI bitmap assets.
    """

    out = FIGURE_DIR / filename
    if out.exists() and out.stat().st_size > 20_000:
        return out
    url = "https://image.pollinations.ai/prompt/" + requests.utils.quote(prompt)
    params = {
        "width": str(width),
        "height": str(height),
        "nologo": "true",
        "enhance": "true",
        "model": "flux",
        "seed": str(abs(hash(filename)) % 1_000_000),
    }
    response = requests.get(url, params=params, timeout=90)
    response.raise_for_status()
    out.write_bytes(response.content)
    with Image.open(out) as img:
        img.convert("RGB").resize((width, height), Image.Resampling.LANCZOS).save(out)
    return out


def make_fallback_concept_image(path: Path, title: str, blocks: list[tuple[str, str, str]]) -> None:
    width, height = 1536, 864
    img = Image.new("RGBA", (width, height), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    title_f = font(54, True)
    h_f = font(34, True)
    b_f = font(25)
    palette = ["#DBEAFE", "#CCFBF1", "#FEF3C7", "#FFE4E6"]
    outline = ["#2563EB", "#0F766E", "#D97706", "#E11D48"]
    draw.text((78, 56), title, fill="#0F172A", font=title_f)
    x0, y0 = 86, 205
    card_w, card_h = 305, 430
    for i, (head, body, icon) in enumerate(blocks):
        x = x0 + i * 355
        add_shadow(img, (x, y0, x + card_w, y0 + card_h), radius=28)
        rounded(draw, (x, y0, x + card_w, y0 + card_h), 28, palette[i % len(palette)], outline[i % len(outline)], 3)
        draw.ellipse((x + 28, y0 + 32, x + 108, y0 + 112), fill=outline[i % len(outline)])
        draw.text((x + 50, y0 + 48), icon, fill="white", font=font(38, True))
        draw.text((x + 28, y0 + 140), wrap_text(head, 16), fill="#111827", font=h_f)
        draw.multiline_text((x + 28, y0 + 238), wrap_text(body, 24), fill="#334155", font=b_f, spacing=10)
        if i < len(blocks) - 1:
            ax = x + card_w + 20
            ay = y0 + card_h // 2
            draw.line((ax, ay, ax + 38, ay), fill="#64748B", width=5)
            draw.polygon([(ax + 38, ay - 14), (ax + 64, ay), (ax + 38, ay + 14)], fill="#64748B")
    img.convert("RGB").save(path)


def draw_network_icon(draw: ImageDraw.ImageDraw, cx: int, cy: int, scale: float, color: str) -> None:
    points = [
        (cx - int(120 * scale), cy - int(45 * scale)),
        (cx - int(38 * scale), cy + int(55 * scale)),
        (cx + int(66 * scale), cy - int(20 * scale)),
        (cx + int(130 * scale), cy + int(72 * scale)),
        (cx + int(15 * scale), cy + int(120 * scale)),
    ]
    for i, p1 in enumerate(points):
        for p2 in points[i + 1 :]:
            if (i + len(points)) % 2 == 0 or abs(i - points.index(p2)) <= 2:
                draw.line((p1[0], p1[1], p2[0], p2[1]), fill="#94A3B8", width=max(2, int(3 * scale)))
    for x, y in points:
        draw.ellipse(
            (x - int(18 * scale), y - int(18 * scale), x + int(18 * scale), y + int(18 * scale)),
            fill=color,
            outline="white",
            width=max(2, int(3 * scale)),
        )
    for r in [50, 82]:
        draw.arc(
            (cx - int(r * scale), cy - int(r * scale), cx + int(r * scale), cy + int(r * scale)),
            210,
            330,
            fill=color,
            width=max(2, int(3 * scale)),
        )


def make_problem_formulation_visual(path: Path) -> None:
    width, height = 1536, 864
    img = Image.new("RGBA", (width, height), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    title_f = font(48, True)
    h_f = font(30, True)
    b_f = font(22)
    draw.text((70, 50), "问题 Formulation：状态 + 动作 -> 未来状态", fill="#0F172A", font=title_f)
    draw.text((72, 112), "目标是学习近似状态转移，聚焦系统动态耦合关系。", fill="#475569", font=b_f)

    panels = [
        (82, 205, 430, 655, "#DBEAFE", "#2563EB", "当前状态", "节点位置 / 链路 / 队列 / 资源"),
        (594, 205, 942, 655, "#CCFBF1", "#0F766E", "调度动作", "卸载 / RB / CPU / 回传"),
        (1106, 205, 1454, 655, "#FEF3C7", "#D97706", "未来演化", "node / link / rate / task"),
    ]
    for x1, y1, x2, y2, fill, stroke, head, body in panels:
        add_shadow(img, (x1, y1, x2, y2), radius=28)
        rounded(draw, (x1, y1, x2, y2), 28, fill, stroke, 3)
        draw.text((x1 + 32, y1 + 30), head, fill="#111827", font=h_f)
        draw.text((x1 + 32, y1 + 78), body, fill="#334155", font=b_f)
        draw_network_icon(draw, (x1 + x2) // 2, y1 + 265, 0.9, stroke)
    for x in [470, 982]:
        draw.line((x, 430, x + 80, 430), fill="#64748B", width=7)
        draw.polygon([(x + 80, 410), (x + 118, 430), (x + 80, 450)], fill="#64748B")

    rounded(draw, (290, 710, 1246, 790), 24, "#FFFFFF", "#CBD5E1", 2)
    draw.text((330, 733), "核心输出：预测未来节点状态、链路活跃性、活跃链路速率和任务状态", fill="#0F172A", font=b_f)
    img.convert("RGB").save(path)


def make_method_pipeline_visual(path: Path) -> None:
    width, height = 1536, 864
    img = Image.new("RGBA", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_f = font(48, True)
    h_f = font(28, True)
    b_f = font(21)
    draw.text((70, 50), "方法链路：从仿真日志到 PI-JWM 诊断", fill="#0F172A", font=title_f)
    stages = [
        ("仿真日志", "状态、任务、链路、动作", "#DBEAFE", "#2563EB"),
        ("数据接口", "dataset / seed\naction interface", "#CCFBF1", "#0F766E"),
        ("模型训练", "动作条件双图 rollout", "#EDE9FE", "#7C3AED"),
        ("指标诊断", "F1 / RMSE / 稳健性", "#FEF3C7", "#D97706"),
    ]
    x0, y0 = 78, 220
    card_w, card_h, gap = 300, 390, 70
    for i, (head, body, fill, stroke) in enumerate(stages):
        x = x0 + i * (card_w + gap)
        add_shadow(img, (x, y0, x + card_w, y0 + card_h), radius=26)
        rounded(draw, (x, y0, x + card_w, y0 + card_h), 26, fill, stroke, 3)
        draw.text((x + 28, y0 + 32), head, fill="#111827", font=h_f)
        draw.multiline_text((x + 28, y0 + 82), wrap_text(body, 26), fill="#334155", font=b_f, spacing=9)
        if i == 0:
            for k in range(4):
                rounded(draw, (x + 74, y0 + 180 + k * 35, x + 226, y0 + 205 + k * 35), 7, "#FFFFFF", stroke, 2)
        elif i == 1:
            for k in range(4):
                draw.rectangle((x + 72 + k * 34, y0 + 190, x + 100 + k * 34, y0 + 320), fill="#FFFFFF", outline=stroke, width=2)
        elif i == 2:
            draw_network_icon(draw, x + 150, y0 + 255, 0.65, stroke)
        else:
            baseline = y0 + 315
            draw.line((x + 56, baseline, x + 240, baseline), fill="#94A3B8", width=3)
            for k, val in enumerate([82, 120, 58, 150]):
                rounded(draw, (x + 70 + k * 40, baseline - val, x + 100 + k * 40, baseline), 6, stroke, None, 1)
        if i < len(stages) - 1:
            ax = x + card_w + 20
            ay = y0 + card_h // 2
            draw.line((ax, ay, ax + 36, ay), fill="#64748B", width=6)
            draw.polygon([(ax + 36, ay - 16), (ax + 66, ay), (ax + 36, ay + 16)], fill="#64748B")
    rounded(draw, (210, 700, 1326, 782), 24, "#F8FAFC", "#CBD5E1", 2)
    draw.text((250, 724), "本阶段重点：先让世界模型稳定预测状态，再接入候选动作评估。", fill="#0F172A", font=b_f)
    img.convert("RGB").save(path)


def make_dual_graph_visual(path: Path) -> None:
    width, height = 1536, 864
    img = Image.new("RGBA", (width, height), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    title_f = font(48, True)
    h_f = font(27, True)
    b_f = font(21)
    draw.text((70, 50), "PI-JWM：物理图 + 信息图联合 rollout", fill="#0F172A", font=title_f)
    draw.text((72, 112), "物理图负责空间运动关系，信息图负责通信、任务与动作历史。", fill="#475569", font=b_f)

    rounded(draw, (80, 195, 555, 395), 28, "#DBEAFE", "#2563EB", 3)
    draw.text((112, 224), "物理图", fill="#111827", font=h_f)
    draw.text((112, 268), "位置 / 距离 / 覆盖 / 相对运动", fill="#334155", font=b_f)
    draw_network_icon(draw, 350, 320, 0.56, "#2563EB")

    rounded(draw, (80, 470, 555, 670), 28, "#CCFBF1", "#0F766E", 3)
    draw.text((112, 498), "信息图", fill="#111827", font=h_f)
    draw.text((112, 542), "任务 / 链路 / 资源 / 动作历史", fill="#334155", font=b_f)
    draw_network_icon(draw, 350, 595, 0.56, "#0F766E")

    add_shadow(img, (720, 300, 1045, 565), radius=34)
    rounded(draw, (720, 300, 1045, 565), 34, "#FFFFFF", "#7C3AED", 4)
    draw.text((768, 342), "联合世界模型", fill="#111827", font=h_f)
    draw.text((768, 392), "双图编码 + 融合 + rollout", fill="#334155", font=b_f)
    for y in [295, 575]:
        draw.line((555, y + 20, 710, 410), fill="#64748B", width=5)
    draw.line((1045, 432, 1175, 432), fill="#64748B", width=6)
    draw.polygon([(1175, 412), (1212, 432), (1175, 452)], fill="#64748B")

    outputs = [("node", "#2563EB"), ("link", "#0F766E"), ("rate", "#D97706"), ("task", "#E11D48")]
    for i, (label, color) in enumerate(outputs):
        y = 245 + i * 115
        rounded(draw, (1220, y, 1450, y + 78), 18, "#FFFFFF", color, 3)
        draw.text((1250, y + 21), label, fill="#111827", font=h_f)
    rounded(draw, (265, 735, 1270, 805), 24, "#FFFFFF", "#CBD5E1", 2)
    draw.text((305, 756), "当前结论：dual 在 active-rate RMSE 与 link-rate RMSE 上最优，双图联合对链路速率更有价值。", fill="#0F172A", font=b_f)
    img.convert("RGB").save(path)


def ensure_ai_figures() -> dict[str, Path]:
    paths: dict[str, Path] = {}
    cover_path = FIGURE_DIR / "cover_pijwm_scene.png"
    if not cover_path.exists() or cover_path.stat().st_size <= 20_000:
        try:
            cover_path = fetch_ai_image("cover_pijwm_scene.png", AI_FIGURES["cover_pijwm_scene.png"])
        except Exception:
            make_fallback_concept_image(
                cover_path,
                "PI-JWM Research Scene",
                [
                    ("Physical Layer", "mobility, distance, coverage, topology", "P"),
                    ("Information Layer", "tasks, links, queues, actions", "I"),
                    ("World Model", "state transition and rollout prediction", "W"),
                    ("Decision Support", "evaluate candidate actions later", "D"),
                ],
            )
    paths["cover_pijwm_scene.png"] = cover_path
    paths["problem_formulation_ai.png"] = FIGURE_DIR / "problem_formulation_ai.png"
    paths["method_pipeline_ai.png"] = FIGURE_DIR / "method_pipeline_ai.png"
    paths["dual_graph_model_ai.png"] = FIGURE_DIR / "dual_graph_model_ai.png"
    make_problem_formulation_visual(paths["problem_formulation_ai.png"])
    make_method_pipeline_visual(paths["method_pipeline_ai.png"])
    make_dual_graph_visual(paths["dual_graph_model_ai.png"])
    if not paths["cover_pijwm_scene.png"].exists():
        make_fallback_concept_image(
            paths["cover_pijwm_scene.png"],
            "PI-JWM Research Scene",
            [
                ("Physical Layer", "mobility, distance, coverage, topology", "P"),
                ("Information Layer", "tasks, links, queues, actions", "I"),
                ("World Model", "state transition and rollout prediction", "W"),
                ("Decision Support", "evaluate candidate actions later", "D"),
            ],
        )
    if not paths["problem_formulation_ai.png"].exists():
        make_fallback_concept_image(
            paths["problem_formulation_ai.png"],
            "Problem Formulation",
            [
                ("Current State", "nodes, links, queues and resources", "S"),
                ("Action", "offloading, RB, CPU and scheduling", "A"),
                ("Future State", "node, link and task evolution", "T"),
                ("Objective", "support stable online decisions", "O"),
            ],
        )
    if not paths["method_pipeline_ai.png"].exists():
        make_fallback_concept_image(
            paths["method_pipeline_ai.png"],
            "Method Pipeline",
            [
                ("Simulation Logs", "reference simulator produces traces", "L"),
                ("Dataset", "state, action and edge-level samples", "D"),
                ("World Model", "action-conditioned dual-graph rollout", "M"),
                ("Diagnostics", "RMSE, F1, robustness and uncertainty", "E"),
            ],
        )
    if not paths["dual_graph_model_ai.png"].exists():
        make_fallback_concept_image(
            paths["dual_graph_model_ai.png"],
            "Dual Graph Model",
            [
                ("Physical Graph", "geometry and movement relations", "P"),
                ("Information Graph", "communication, task and action history", "I"),
                ("Fusion", "joint message passing and latent rollout", "F"),
                ("Heads", "node, link, rate and task prediction", "H"),
            ],
        )
    return paths


def make_metric_chart(summary: dict, path: Path) -> None:
    runs = summary["real_data_sanity"]["runs"]
    modes = ["dual", "physical_only", "information_only"]
    mode_labels = {"dual": "dual", "physical_only": "physical", "information_only": "information"}
    colors = {"dual": "#2563EB", "physical_only": "#10B981", "information_only": "#F97316"}
    metrics = [
        ("active-rate RMSE", [runs[m]["test_eval"]["active_rate"]["active_rmse"] for m in modes], 0),
        ("link-rate RMSE", [runs[m]["test_eval"]["link_rate"]["rmse"] for m in modes], 1),
        ("task RMSE", [runs[m]["test_eval"]["task"]["rmse"] for m in modes], 2),
        ("node RMSE", [runs[m]["test_eval"]["node"]["rmse"] for m in modes], 3),
    ]

    width, height = 1750, 1080
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_f = font(52, True)
    subtitle_f = font(25)
    axis_f = font(23, True)
    small_f = font(21)
    draw.text((70, 48), "PI-JWM v6 full80：三种图模式测试集指标对比", fill="#0F172A", font=title_f)
    draw.text(
        (72, 120),
        "activity F1 均为 1.0；主要差异体现在 active-rate、link-rate、task 与 node 的原尺度误差。",
        fill="#475569",
        font=subtitle_f,
    )

    panel_w, panel_h = 770, 340
    positions = [(80, 205), (900, 205), (80, 610), (900, 610)]
    for (metric_name, values, metric_idx), (px, py) in zip(metrics, positions):
        rounded(draw, (px, py, px + panel_w, py + panel_h), 20, "#F8FAFC", "#CBD5E1", 2)
        draw.text((px + 28, py + 22), metric_name, fill="#111827", font=axis_f)
        draw.text((px + 28, py + 52), "越低越好", fill="#64748B", font=small_f)
        max_v = max(values) * 1.15
        min_v = min(values)
        bar_area_x = px + 60
        bar_base_y = py + 260
        bar_max_h = 150
        gap = 180
        for i, (mode, value) in enumerate(zip(modes, values)):
            x = bar_area_x + i * gap
            h = value / max_v * bar_max_h
            y = bar_base_y - h
            rounded(draw, (x, y, x + 90, bar_base_y), 12, colors[mode], None, 1)
            value_color = "#1D4ED8" if math.isclose(value, min_v) else "#111827"
            draw.text((x - 20, y - 34), f"{value:.2f}", fill=value_color, font=small_f)
            draw.text((x - 18, bar_base_y + 18), mode_labels[mode], fill="#334155", font=small_f)
        draw.line((bar_area_x - 20, bar_base_y, bar_area_x + 450, bar_base_y), fill="#94A3B8", width=2)
    img.save(path)


def make_small_table_chart(summary: dict, path: Path) -> None:
    runs = summary["real_data_sanity"]["runs"]
    rows = [
        ("dual", 19, 0.69, 1.0, 228.318431, 6.416102, 38.282549, 3.664203),
        ("physical_only", 28, 0.77, 1.0, 268.036374, 7.473947, 37.067164, 4.855683),
        ("information_only", 19, 0.45, 1.0, 235.526004, 6.563317, 37.305306, 3.380836),
    ]
    headers = ["模式", "epoch", "阈值", "F1", "active-rate", "link-rate", "node", "task"]
    width, height = 1650, 520
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_f = font(42, True)
    cell_f = font(24)
    head_f = font(25, True)
    draw.text((60, 38), "v6 full80 测试集核心数值", fill="#0F172A", font=title_f)
    x0, y0 = 60, 130
    col_w = [250, 150, 150, 130, 240, 220, 190, 180]
    row_h = 70
    rounded(draw, (x0 - 10, y0 - 10, x0 + sum(col_w) + 10, y0 + row_h * 4 + 10), 18, "#F8FAFC", "#CBD5E1", 2)
    x = x0
    for i, h in enumerate(headers):
        draw.rectangle((x, y0, x + col_w[i], y0 + row_h), fill="#DBEAFE")
        draw.text((x + 14, y0 + 20), h, fill="#111827", font=head_f)
        x += col_w[i]
    for r, row in enumerate(rows, start=1):
        y = y0 + row_h * r
        x = x0
        for i, value in enumerate(row):
            fill = "#FFFFFF" if r % 2 else "#F1F5F9"
            draw.rectangle((x, y, x + col_w[i], y + row_h), fill=fill)
            txt = f"{value:.6f}" if isinstance(value, float) and i >= 4 else str(value)
            if i == 2:
                txt = f"{value:.2f}"
            if i == 3:
                txt = f"{value:.1f}"
            color = "#1D4ED8" if (row[0] == "dual" and i in [4, 5]) else "#111827"
            draw.text((x + 14, y + 20), txt, fill=color, font=cell_f)
            x += col_w[i]
    img.save(path)


def set_doc_font(doc: Document) -> None:
    styles = doc.styles
    for style_name in ["Normal", "Body Text", "List Bullet", "List Number"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "宋体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(10.5)
    heading_specs = {
        "Title": (20, "微软雅黑"),
        "Heading 1": (15, "微软雅黑"),
        "Heading 2": (12.5, "微软雅黑"),
        "Heading 3": (11.5, "微软雅黑"),
    }
    for style_name, (size, name) in heading_specs.items():
        style = styles[style_name]
        style.font.name = name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(15, 23, 42)


def set_paragraph_format(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.style.name == "Normal":
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.first_line_indent = Pt(21)
        elif p.style.name.startswith("Heading"):
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CBD5E1")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_font(cell, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(9.5)
            run.bold = bold


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    add_bottom_border(p)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(71, 85, 105)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(0)


def add_picture(doc: Document, path: Path, width_cm: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "DBEAFE")
        set_cell_font(cell, bold=True)
        if widths:
            cell.width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_font(cells[i])
            if widths:
                cells[i].width = Cm(widths[i])


def metric_rows() -> list[list[str]]:
    return [
        ["dual", "19", "0.69", "1.000000", "228.318431", "6.416102", "38.282549", "3.664203"],
        ["physical_only", "28", "0.77", "1.000000", "268.036374", "7.473947", "37.067164", "4.855683"],
        ["information_only", "19", "0.45", "1.000000", "235.526004", "6.563317", "37.305306", "3.380836"],
    ]


def build_docx(figures: dict[str, Path], metric_chart: Path, metric_table_chart: Path) -> None:
    doc = Document()
    set_doc_font(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    run = title.add_run("科研实习进组以来工作总结")
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(15, 23, 42)
    subtitle = doc.add_paragraph("禹尧珅｜2026 年 5 月 31 日线上讨论｜PI-JWM 阶段进展")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = Pt(0)
    subtitle.runs[0].font.name = "微软雅黑"
    subtitle.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(71, 85, 105)

    add_picture(
        doc,
        figures["cover_pijwm_scene.png"],
        16.2,
        "图 1  研究场景示意：低空/车联网节点、边缘服务器、无线链路与任务流共同构成物理-信息耦合系统。",
    )

    doc.add_paragraph(
        "本次汇报围绕老师提出的四个问题展开：问题 formulation、用了什么方法、克服了什么困难、有什么创新。"
        "从进组以来的工作看，我当前的主线已经收敛为 PI-JWM，即 Physical-Information Joint World Model。"
        "它的目标是在网络化具身智能体场景中，联合建模物理网络和信息网络的状态演化，为后续通信资源、计算卸载和运动轨迹协同决策提供可预测的系统模型。"
    )

    add_h1(doc, "一、问题 Formulation")
    doc.add_paragraph(
        "老师最初给出的方向可以概括为：在基站覆盖下的联网具身智能体系统中，多个无人机、车辆或边缘节点需要同时完成感知、通信、计算和移动。"
        "系统面对的是会随动作持续变化的动态网络，需要联合预测节点、链路、任务和资源状态。"
        "节点位置会改变可通信拓扑，调度动作会改变任务队列和资源占用，任务到达和信道衰减又会带来随机扰动。"
    )
    add_picture(
        doc,
        figures["problem_formulation_ai.png"],
        15.6,
        "图 2  问题定义示意：当前状态、调度动作与未来状态之间存在连续耦合关系。",
    )
    add_bullets(
        doc,
        [
            "输入状态：节点位置/速度、链路活动、链路速率、任务队列、资源占用、历史动作等。",
            "动作变量：任务卸载、RB/CPU 分配、回传调度、边级通信动作，后续可进一步接入轨迹动作。",
            "预测目标：未来节点状态、链路是否活跃、活跃链路速率、任务状态和资源变化。",
            "最终用途：先得到稳定可靠的世界模型，再让它评估候选调度或轨迹动作，为在线决策提供近似环境反馈。",
        ],
    )

    add_h1(doc, "二、用了什么方法")
    doc.add_paragraph(
        "方法路线分为三层。第一层是仿真与数据层，用 AirFogSim/SUMO 作为参考仿真器和数据来源，解析单步推进机制并导出可训练样本。"
        "第二层是预测模型层，从 state-only、state-action、edge-level、two-stage 等基线逐步推进到 world model。"
        "第三层是 PI-JWM 的双图建模层，将物理图和信息图分开编码，再联合预测未来状态。"
    )
    add_picture(
        doc,
        figures["method_pipeline_ai.png"],
        15.6,
        "图 3  方法链路示意：仿真日志生成样本，样本训练世界模型，世界模型输出诊断指标和后续决策依据。",
    )

    doc.add_heading("1. 数据与动作接口", level=2)
    add_bullets(
        doc,
        [
            "dataset：把仿真日志转成监督学习样本，包含当前状态和下一步真实状态。",
            "multi-seed：用不同随机种子生成多组轨迹，检查模型是否只记住某一次随机过程。",
            "strict action：保留调度动作的严格定义，避免把动作信息混进状态标签造成信息泄漏。",
            "edge action：把全局调度动作映射到候选链路层面，使链路预测能看到与该边相关的动作信息。",
        ],
    )

    doc.add_heading("2. 模型演进", level=2)
    add_table(
        doc,
        ["阶段", "核心思想", "作用"],
        [
            ["普通预测模型", "只用状态或状态+动作做下一步预测", "建立可训练基线，判断动作信息是否有用。"],
            ["边级链路模型", "把链路活动与链路速率单独建模", "定位通信链路侧的误差来源。"],
            ["两阶段模型", "先判断是否活跃，再预测活跃链路的 rate", "缓解大量非活跃链路对速率回归的干扰。"],
            ["World Model v0-v3", "统一输入输出、分阶段训练、latent rollout、通信图消息传递", "从一步预测走向可滚动预测。"],
            ["PI-JWM v6", "物理图+信息图+动作历史联合建模", "当前主线模型，用于验证双图是否提升状态预测。"],
        ],
        widths=[3.0, 6.6, 6.5],
    )

    add_h1(doc, "三、克服了什么困难")
    add_table(
        doc,
        ["困难", "具体表现", "处理方式"],
        [
            [
                "问题范围过宽",
                "基座模型、世界模型、资源调度、轨迹控制和网络预测都相关，容易没有主线。",
                "将阶段目标收敛到 PI-JWM：先做好物理-信息联合状态预测，再考虑决策接口。",
            ],
            [
                "仿真机制不透明",
                "任务、信道、动作、队列和资源更新的顺序需要从代码与日志中梳理。",
                "解析仿真单步推进，把 AirFogSim 固定为数据来源和机制参照。",
            ],
            [
                "动作与链路难对齐",
                "全局动作不能直接解释某条候选边未来是否活跃、速率是多少。",
                "构建 strict action 与 edge action，把动作信息映射到边级样本。",
            ],
            [
                "链路标签稀疏",
                "多数候选边不活跃，直接预测 rate 会被大量零值影响。",
                "拆成 activity prediction 与 active-rate regression，并分别报告 F1 和 RMSE。",
            ],
            [
                "阶段方向容易偏到排序接口",
                "v5 selector/ranking 适合做决策接口诊断，主模型仍回到状态预测。",
                "明确 v5 是决策接口诊断，v6 回到双图世界模型主线。",
            ],
        ],
        widths=[3.0, 6.1, 6.7],
    )

    add_h1(doc, "四、有什么创新")
    add_picture(
        doc,
        figures["dual_graph_model_ai.png"],
        15.6,
        "图 4  PI-JWM 双图建模示意：物理图表达空间运动关系，信息图表达通信、任务与动作历史，二者联合进行 rollout。",
    )
    add_bullets(
        doc,
        [
            "问题创新：将研究对象明确为物理网络和信息网络联合演化，而非孤立的链路、任务或节点预测。",
            "表示创新：把物理图与信息图拆开建模，分别承载几何移动关系和通信任务关系，再在预测层融合。",
            "动作创新：将调度动作从日志中抽取出来，并进一步构造成边级动作，使模型具备动作条件预测能力。",
            "评估创新：同时看 activity F1、active-rate RMSE、link-rate RMSE、node RMSE、task RMSE，避免只依赖一个总 loss。",
            "应用潜力：当世界模型足够稳定后，可以用它快速评估候选卸载、资源分配与轨迹动作，减少直接仿真搜索成本。",
        ],
    )

    add_h1(doc, "五、当前实验结果")
    doc.add_paragraph(
        "目前最新的 v6 full80 GPU 实验使用 train=1520、val=190、test=190 的真实仿真样本，三种模式均训练 80 epoch。"
        "对比对象包括 dual、physical_only 和 information_only。这里的 dual 表示同时使用物理图和信息图；physical_only 只保留几何与移动关系；information_only 只保留通信、任务与动作历史关系。"
    )
    add_picture(
        doc,
        metric_chart,
        16.2,
        "图 5  数据图：v6 full80 测试集指标对比。除 activity F1 外，其余 RMSE 均为越低越好。",
    )
    add_picture(
        doc,
        metric_table_chart,
        16.0,
        "图 6  数据表图：v6 full80 测试集原始数值，便于汇报时直接指读。",
    )
    add_table(
        doc,
        ["模式", "best epoch", "threshold", "activity F1", "active-rate RMSE", "link-rate RMSE", "node RMSE", "task RMSE"],
        metric_rows(),
        widths=[2.6, 1.7, 1.7, 1.8, 2.5, 2.2, 1.8, 1.8],
    )
    doc.add_paragraph(
        "阶段结论是：三种模式都能稳定判断链路是否活跃，test activity F1 均为 1.0。"
        "dual 在 active-rate RMSE 和 link-rate RMSE 上最好，说明物理图与信息图联合后，对活跃链路速率幅值的预测更有帮助。"
        "information_only 的 task RMSE 最低，说明任务状态变化更依赖信息图历史和动作。"
        "physical_only 的 node RMSE 最低，但链路和任务指标较弱，说明只看几何物理关系不足以解释通信任务演化。"
    )

    add_h1(doc, "六、后续计划")
    add_bullets(
        doc,
        [
            "补普通模型对比：把 MLP、Ridge、state-action predictor、non-rollout predictor 与 PI-JWM v6 放到同一指标体系中比较，回答“世界模型相比普通模型有什么优势”。",
            "改进 active-rate 建模：重点降低活跃链路速率幅值误差，尝试 active-only regression、加权损失、两阶段 head 和分布式/不确定性输出。",
            "增强双图联合建模：继续优化 physical graph 与 information graph 的融合方式，加入跨图注意力或门控，让两类关系在不同预测目标上自适应分工。",
            "做稳健性与泛化实验：补充跨 seed、输入扰动、阈值迁移、置信区间和 per-seed 曲线，验证模型在多次数据与扰动条件下的有效性。",
            "连接决策接口：当状态 rollout 更稳定后，用 PI-JWM 评估候选卸载、资源分配和轨迹动作，再把 v5 selector/ranking 作为诊断接口接回来。",
        ],
    )

    add_h1(doc, "七、15 分钟讲述安排")
    add_table(
        doc,
        ["时间", "内容", "讲述重点"],
        [
            ["0-2 分钟", "背景与 formulation", "为什么需要物理-信息联合世界模型。"],
            ["2-5 分钟", "方法路线", "数据、动作接口、基线模型和 world model 演进。"],
            ["5-8 分钟", "困难与解决", "问题收敛、机制解析、动作对齐、稀疏链路。"],
            ["8-11 分钟", "创新点", "双图表示、动作条件预测、多指标诊断。"],
            ["11-14 分钟", "实验结果", "activity F1、active-rate、link-rate、task/node 的分工解释。"],
            ["14-15 分钟", "下一步计划", "普通模型对比、active-rate 改进、稳健性与决策接口。"],
        ],
        widths=[2.6, 4.0, 9.2],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_h1(doc, "附录：进组以来已完成工作")
    add_bullets(
        doc,
        [
            "完成相关方向阅读与主线收敛：从 ST-GNN、world model、低空网络和联网具身智能体，收敛到 PI-JWM。",
            "完成仿真机制解析：梳理任务产生、信道衰减、动作注入、资源执行和日志导出机制。",
            "完成数据接口：dataset、多 seed、strict action、edge action、world model dataset。",
            "完成多类基线：state-only、state-action、structured、edge-level、two-stage、edge-action baseline。",
            "完成世界模型演进：v0/v1/v2/v3/v4/v6 的阶段实现与诊断。",
            "完成 v6 full80 GPU 三模式消融：dual、physical_only、information_only，并形成当前阶段结论。",
            "完成本地项目结构整理：PI-JWM 代码、脚本、测试、reference 和 artifacts 已按规范放置。",
        ],
    )

    set_paragraph_format(doc)
    doc.save(DOCX_PATH)


def build_speaker_script() -> None:
    content = """# 科研实习进组以来工作总结讲稿

时间：2026 年 5 月 31 日线上讨论<br>
汇报人：禹尧珅<br>
主线：PI-JWM，Physical-Information Joint World Model

## 0. 开场

老师好，各位同学好。我这次汇报主要按照老师通知里的四个问题来讲：第一是问题 formulation，第二是用了什么方法，第三是克服了什么困难，第四是目前有什么创新和后续计划。

我进组以来的主线现在已经收敛为 PI-JWM，也就是物理-信息联合世界模型。这里需要先说明一点：AirFogSim/SUMO 在我的工作里是仿真器和数据来源，用来产生可分析、可训练的网络动态日志；真正要做的方法主线，是我们自己的 PI-JWM。

## 1. 问题 Formulation

我理解的课题背景是，在基站覆盖下的联网具身智能体系统中，无人机、车辆、边缘节点会同时涉及运动、通信、计算和任务处理。这个系统的难点在于动态耦合：节点会移动，链路会波动，任务和资源也会被动作持续改变。

节点位置变化会改变物理拓扑，通信动作和计算卸载动作会改变任务队列、资源占用和未来链路状态，任务到达和信道衰减又带来随机性。所以我们需要学习一个近似的状态转移模型：给定当前状态和动作，预测未来节点、链路和任务会怎么变化。

因此我的 formulation 是：输入当前的物理状态、信息状态和调度动作，输出未来的节点状态、链路活动、链路速率和任务状态。最终希望这个模型能作为后续在线决策的近似环境，用世界模型快速评估候选动作，减少对完整仿真的反复调用。

## 2. 用了什么方法

方法上我分三层推进。

第一层是数据层。我先解析仿真器内部机制，包括任务产生、信道衰减、动作注入、资源执行、队列更新和日志导出。基于这些日志，我构建了 dataset、多 seed 数据、strict action、edge action 和 world model dataset。这里的重点是让样本里有明确的“状态、动作、下一步状态”关系。

第二层是基线层。我先做 state-only 模型，只看历史状态；再做 state-action 模型，把动作加入输入；然后做 edge-level 和 two-stage link model，把链路是否活跃和活跃链路速率拆开。这样可以定位误差到底来自节点、任务，还是链路速率。

第三层是 PI-JWM 双图世界模型。物理图主要表达位置、距离、覆盖、相对运动这些几何关系；信息图表达任务、链路、队列、资源、动作历史这些信息关系。v6 的核心就是把这两张图联合起来做 rollout，预测未来状态。

## 3. 克服了什么困难

第一个困难是问题范围太宽。最开始方向里有基座模型、世界模型、资源调度、轨迹控制和网络预测，容易发散。后来我把阶段目标收敛成 PI-JWM：先把物理-信息联合状态预测做好，再谈决策接口。

第二个困难是仿真机制不透明。日志里的数据来自任务到达、信道衰减、动作执行和资源更新顺序。我做了代码和日志级别的梳理，把仿真器固定为数据来源，主线保持在 PI-JWM。

第三个困难是动作和链路难对齐。全局动作不能直接解释某条候选边未来是否活跃，所以我做了 strict action 和 edge action，把动作信息映射到边级样本。

第四个困难是链路标签稀疏。多数候选链路其实是不活跃的，如果直接预测 rate，模型很容易被大量零值主导。为了解决这个问题，我把它拆成 activity prediction 和 active-rate regression，分别看 F1 和 RMSE。

第五个困难是阶段方向容易偏到 selector/ranking。v5 确实能做决策接口诊断，所以现在我把 v5 定位成后续诊断接口，把主线拉回 v6 双图世界模型。

## 4. 有什么创新

我认为当前阶段主要有五点创新。

第一是问题层面的创新：问题被定义为物理网络和信息网络的联合演化建模，覆盖链路、任务和节点状态。

第二是表示层面的创新：物理图和信息图分开建模。物理图处理几何移动关系，信息图处理通信、任务和动作历史，再在世界模型里融合。

第三是动作层面的创新：从仿真日志中显式抽取动作，并构造边级动作，让模型具备动作条件预测能力。

第四是评估层面的创新：分别看 activity F1、active-rate RMSE、link-rate RMSE、node RMSE、task RMSE，避免只依赖一个总 loss。这样能看出不同图结构对不同预测目标的贡献。

第五是应用潜力：如果后续世界模型稳定，可以用它快速评估候选卸载、资源分配和轨迹动作，为在线协同决策服务。

## 5. 当前实验结果

目前最新结果是 v6 full80 GPU 实验。数据划分是 train=1520、val=190、test=190，三种模式都训练 80 epoch。

三种模式分别是 dual、physical_only 和 information_only。dual 同时用物理图和信息图；physical_only 只看几何和移动关系；information_only 只看通信、任务和动作历史。

结果上，三种模式的 activity F1 都是 1.0，说明链路是否活跃已经可以稳定学到。真正拉开差距的是速率幅值和任务/节点状态。

dual 的 active-rate RMSE 是 228.318，link-rate RMSE 是 6.416，这两项都是三种模式里最低的。说明物理图和信息图联合后，对活跃链路速率和整体链路速率预测更有帮助。

information_only 的 task RMSE 是 3.381，是任务指标里最低的。这个结果说明任务演化更依赖信息图历史和动作。

physical_only 的 node RMSE 是 37.067，是节点指标里最低的，但它的链路和任务指标较弱。这个结果说明几何物理关系对节点状态有帮助，但单独看物理图不足以解释通信和任务演化。

所以当前阶段结论是：PI-JWM v6 已经能跑通真实数据训练、测试和三模式消融，并且初步验证了双图联合对链路速率预测有价值。

## 6. 后续计划

下一步我准备做五件事。

第一，补普通模型和世界模型的统一对比。也就是把 MLP、Ridge、state-action predictor、non-rollout predictor 和 PI-JWM v6 放到同一指标体系里，回答学长问的“有没有做世界模型和普通模型比较”。

第二，继续改 active-rate 建模。现在链路是否活跃已经学得比较稳，主要误差来自活跃链路速率幅值。后续会尝试 active-only regression、加权损失、两阶段 head 和不确定性输出。

第三，增强双图联合方式。现在是初步 dual graph，后续可以加入跨图注意力或门控，让物理图和信息图在不同预测目标上自适应分工。

第四，补稳健性和泛化实验，包括跨 seed、输入扰动、阈值迁移、置信区间和 per-seed 曲线。

第五，等状态 rollout 更稳定后，再接回决策接口，用世界模型评估候选卸载、资源分配和轨迹动作。

## 7. 结尾

总结一下，我这段时间主要完成了从问题理解、仿真机制解析、数据构建、动作条件预测，到 PI-JWM v6 双图世界模型的阶段推进。当前结果证明这个方向已经能形成可训练、可评估、可解释的实验链路。后续重点是补普通模型对比、降低 active-rate 误差，并把双图世界模型进一步推向可用于决策的稳定 rollout。
"""
    SCRIPT_PATH.write_text(content, encoding="utf-8")


def update_plan_file() -> None:
    plan_path = WORKSPACE_ROOT / "本地计划表.md"
    try:
        text = plan_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = plan_path.read_text(encoding="gbk", errors="ignore")
    marker = "## 2026-05-30 5.31 线上讨论材料 v2"
    addition = f"""

{marker}

- 已按老师的临时讨论要求重做屏幕共享版 DOCX：`meeting/5.31/{DOCX_PATH.name}`。
- 已新增配套讲稿：`meeting/5.31/{SCRIPT_PATH.name}`。
- 材料主线统一为 PI-JWM，AirFogSim/SUMO 仅作为参考仿真器和数据来源。
- 非数据概念图已重新生成/重绘；数据相关图保留 v6 full80 的精确指标。
"""
    if marker not in text:
        plan_path.write_text(text.rstrip() + addition + "\n", encoding="utf-8")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    probe = FIGURE_DIR / "_pollinations_probe.png"
    if probe.exists():
        probe.unlink()
    summary = json.loads(DATA_SUMMARY_PATH.read_text(encoding="utf-8"))
    ai_figures = ensure_ai_figures()
    metric_chart = FIGURE_DIR / "v6_full80_metrics_v2.png"
    metric_table_chart = FIGURE_DIR / "v6_full80_metric_table_v2.png"
    make_metric_chart(summary, metric_chart)
    make_small_table_chart(summary, metric_table_chart)
    build_docx(ai_figures, metric_chart, metric_table_chart)
    build_speaker_script()
    update_plan_file()
    print(DOCX_PATH)
    print(SCRIPT_PATH)


if __name__ == "__main__":
    main()
